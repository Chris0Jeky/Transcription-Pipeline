"""
qa_format.py
------------
Turns diarized transcription segments into a clean, readable transcript.

Input: list of segments, each: {"start","end","speaker","text"}
(the shape WhisperX produces after diarization)

Two label styles (set qa_format.LABEL_STYLE):
  "P"  -> P1, P2, P3 ... one per distinct speaker (order of first appearance)
  "QA" -> Question / Answer (interviewer auto-detected from question phrasing)

Outputs: write_docx(), write_txt(), write_srt(); process() does all three.
"""
import re

LABEL_STYLE = "P"   # "P" or "QA"

QUESTION_WORDS = ("what", "why", "how", "do you", "did you", "have you",
                  "would you", "could you", "can you", "is there", "are you",
                  "tell me", "tell us", "what's", "whats", "who", "when",
                  "where", "which", "anyone", "does anyone")


def _is_question(text):
    t = text.strip().lower()
    if "?" in t:
        return True
    return any(t.startswith(w) or (" " + w) in t[:40] for w in QUESTION_WORDS)


def merge_turns(segments, max_gap=1.5):
    turns = []
    for s in segments:
        spk = s.get("speaker", "UNKNOWN")
        txt = (s.get("text") or "").strip()
        if not txt:
            continue
        if turns and turns[-1]["speaker"] == spk and \
           (s["start"] - turns[-1]["end"]) <= max_gap:
            turns[-1]["text"] += " " + txt
            turns[-1]["end"] = s["end"]
        else:
            turns.append({"speaker": spk, "text": txt,
                          "start": s["start"], "end": s["end"]})
    return turns


def detect_interviewer(turns):
    stats = {}
    for t in turns:
        d = stats.setdefault(t["speaker"], {"turns": 0, "q": 0, "words": 0})
        d["turns"] += 1
        d["words"] += len(t["text"].split())
        if _is_question(t["text"]):
            d["q"] += 1
    if not stats:
        return None
    total_words = sum(d["words"] for d in stats.values()) or 1
    best, best_score = None, -1
    for spk, d in stats.items():
        q_ratio = d["q"] / d["turns"] if d["turns"] else 0
        talk_share = d["words"] / total_words
        score = q_ratio - 0.3 * talk_share
        if score > best_score:
            best_score, best = score, spk
    return best


def assign_roles(turns, interviewer_speaker=None):
    """Attach 'role' to each turn according to LABEL_STYLE. Returns
    (turns, speaker->label map, n_speakers)."""
    order = list(dict.fromkeys(t["speaker"] for t in turns))
    if LABEL_STYLE == "P":
        label_map = {spk: "P" + str(i + 1) for i, spk in enumerate(order)}
        for t in turns:
            t["role"] = label_map[t["speaker"]]
        return turns, label_map, len(order)
    # QA style
    if interviewer_speaker is None:
        interviewer_speaker = detect_interviewer(turns)
    answerers = [s for s in order if s != interviewer_speaker]
    multi = len(answerers) > 1
    ans_map = {spk: i + 1 for i, spk in enumerate(answerers)}
    label_map = {}
    for t in turns:
        if t["speaker"] == interviewer_speaker:
            t["role"] = "Question"; label_map[t["speaker"]] = "Question"
        else:
            t["role"] = "Answer (" + str(ans_map[t["speaker"]]) + ")" if multi else "Answer"
            label_map[t["speaker"]] = t["role"]
    return turns, label_map, len(order)


def _fmt_ts(seconds, srt=False):
    ms = int(round((seconds - int(seconds)) * 1000))
    s = int(seconds); h, s = divmod(s, 3600); m, s = divmod(s, 60)
    if srt:
        return "%02d:%02d:%02d,%03d" % (h, m, s, ms)
    return "%02d:%02d" % (m + h * 60, s)


def write_txt(turns, path, with_timestamps=True):
    lines = []
    for t in turns:
        ts = "[" + _fmt_ts(t["start"]) + "] " if with_timestamps else ""
        lines.append(ts + t["role"] + ": " + t["text"])
    open(path, "w", encoding="utf-8").write("\n\n".join(lines) + "\n")


def write_srt(turns, path):
    with open(path, "w", encoding="utf-8") as f:
        for i, t in enumerate(turns, 1):
            f.write(str(i) + "\n" + _fmt_ts(t["start"], True) + " --> "
                    + _fmt_ts(t["end"], True) + "\n" + t["role"] + ": "
                    + t["text"] + "\n\n")


# a palette of distinct label colours (RGB)
_PALETTE = [(0xB0, 0x30, 0x30), (0x1F, 0x38, 0x64), (0x2E, 0x7D, 0x32),
            (0x6A, 0x1B, 0x9A), (0xC6, 0x6A, 0x00), (0x00, 0x69, 0x8C),
            (0xAD, 0x14, 0x57), (0x45, 0x27, 0x00)]


def write_docx(turns, path, title="Transcript", meta=None):
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    p = doc.add_paragraph(); r = p.add_run(title)
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    if meta:
        m = doc.add_paragraph(); mr = m.add_run(meta)
        mr.italic = True; mr.font.size = Pt(9); mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    labels = list(dict.fromkeys(t["role"] for t in turns))
    colour = {lab: RGBColor(*_PALETTE[i % len(_PALETTE)]) for i, lab in enumerate(labels)}
    for t in turns:
        para = doc.add_paragraph()
        lab = para.add_run(t["role"] + ": ")
        lab.bold = True; lab.font.color.rgb = colour[t["role"]]
        para.add_run(t["text"])
    doc.save(path)


def process(segments, out_stem, title="Transcript", meta="",
            interviewer_speaker=None, timestamps=True):
    turns = merge_turns(segments)
    turns, label_map, n_speakers = assign_roles(turns, interviewer_speaker)
    write_docx(turns, out_stem + ".docx", title=title, meta=meta)
    write_txt(turns, out_stem + ".txt", with_timestamps=timestamps)
    write_srt(turns, out_stem + ".srt")
    return {"turns": len(turns), "n_speakers": n_speakers, "labels": label_map}
