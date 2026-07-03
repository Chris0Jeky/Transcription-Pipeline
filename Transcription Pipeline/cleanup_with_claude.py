#!/usr/bin/env python3
"""
cleanup_with_claude.py
----------------------
Second pass over raw transcripts: sends them to Claude (Anthropic API) to
 - add punctuation & capitalisation,
 - fix obvious mis-transcriptions from context + the glossary (e.g. names),
 - WITHOUT changing meaning or inventing words,
 - FLAG anything it isn't sure about, with the timestamp, so you can listen
   to just those spots and correct them by hand.

For every  <name>.txt  in the input folder it writes:
    <name>.cleaned.docx   tidied transcript
    <name>.cleaned.txt    tidied transcript (plain text)
    <name>.FLAGS.txt      list of timestamps + reasons to check manually

Usage:
    set ANTHROPIC_API_KEY=sk-ant-...           (Windows)
    python cleanup_with_claude.py --input "transcripts"
    python cleanup_with_claude.py --input "transcripts" --model claude-sonnet-5

Get an API key at https://console.anthropic.com  ->  API Keys.
Install the SDK:  pip install anthropic
"""
import argparse, os, glob, re, sys

TS_LINE = re.compile(r"^\[(\d{1,2}:\d{2})\]\s*([^:]{1,40}?):\s*(.*)$")


def load_dotenv():
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(here, ".env"), os.path.join(os.getcwd(), ".env")):
        if os.path.exists(p):
            for line in open(p, encoding="utf-8"):
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    k, v = line.split("=", 1)
                    os.environ.setdefault(k.strip(), v.strip())

SYSTEM = (
    "You are cleaning up an automatically-generated interview transcript for a "
    "youth podcast called 'From the Boys'. You will get numbered turns, each with "
    "a timestamp and a speaker label. For each turn:\n"
    "1. Add natural punctuation, capitalisation and paragraphing.\n"
    "2. Fix CLEAR speech-to-text errors using context and the glossary of names/terms "
    "(e.g. mis-heard names, homophones). Prefer glossary spellings for names.\n"
    "3. Remove obvious duplicated stutters ('the the') but keep the speaker's real words "
    "and meaning. Do NOT paraphrase, summarise, add, or remove content.\n"
    "4. If you are UNSURE about a word/name, keep your best guess but wrap it like "
    "{{guess?}}, and add a flag.\n"
    "Return STRICT JSON only: {\"turns\":[{\"i\":<int>,\"text\":\"<cleaned text>\"}...],"
    "\"flags\":[{\"i\":<int>,\"reason\":\"<short reason>\"}...]}. "
    "Keep the same turn indices you were given. Do not include the timestamp or label in "
    "the text - just the cleaned words."
)


def parse_txt(path):
    turns = []
    for line in open(path, encoding="utf-8"):
        line = line.rstrip("\n")
        m = TS_LINE.match(line)
        if m:
            turns.append({"ts": m.group(1), "label": m.group(2).strip(), "text": m.group(3)})
        elif line.strip() and turns:
            turns[-1]["text"] += " " + line.strip()
    return turns


def call_claude(client, model, glossary, batch):
    import json
    payload = {"glossary": glossary,
               "turns": [{"i": i, "label": t["label"], "text": t["text"]}
                         for i, t in batch]}
    msg = client.messages.create(
        model=model, max_tokens=8000, system=SYSTEM,
        messages=[{"role": "user", "content": json.dumps(payload, ensure_ascii=False)}])
    raw = "".join(b.text for b in msg.content if getattr(b, "type", "") == "text")
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("```", 2)[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw)


def main():
    load_dotenv()
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True, help="folder containing <name>.txt transcripts")
    ap.add_argument("--model", default="claude-sonnet-5")
    ap.add_argument("--glossary", default="glossary.txt")
    ap.add_argument("--batch", type=int, default=25, help="turns per API call")
    ap.add_argument("--api-key", default=os.environ.get("ANTHROPIC_API_KEY"))
    args = ap.parse_args()

    if not args.api_key:
        print("ERROR: set ANTHROPIC_API_KEY (or pass --api-key). See header."); sys.exit(2)
    try:
        import anthropic
    except ImportError:
        print("ERROR: pip install anthropic"); sys.exit(2)
    from docx import Document
    from docx.shared import Pt, RGBColor

    glossary = ""
    if os.path.exists(args.glossary):
        glossary = ", ".join(l.strip() for l in open(args.glossary, encoding="utf-8")
                             if l.strip() and not l.startswith("#"))
    client = anthropic.Anthropic(api_key=args.api_key)

    files = [f for f in sorted(glob.glob(os.path.join(args.input, "*.txt")))
             if not f.endswith(".cleaned.txt") and not f.endswith(".FLAGS.txt")]
    if not files:
        print("No .txt transcripts found in " + args.input); sys.exit(1)

    for path in files:
        name = os.path.splitext(os.path.basename(path))[0]
        turns = parse_txt(path)
        print(name + ": " + str(len(turns)) + " turns -> cleaning...")
        flags = []
        for start in range(0, len(turns), args.batch):
            batch = list(enumerate(turns))[start:start + args.batch]
            try:
                res = call_claude(client, args.model, glossary, batch)
            except Exception as e:
                print("  batch @" + str(start) + " failed: " + str(e) + " (kept raw)")
                continue
            for item in res.get("turns", []):
                i = item.get("i")
                if isinstance(i, int) and 0 <= i < len(turns):
                    turns[i]["text"] = item.get("text", turns[i]["text"])
            for fl in res.get("flags", []):
                i = fl.get("i")
                if isinstance(i, int) and 0 <= i < len(turns):
                    flags.append((turns[i]["ts"], turns[i]["label"], fl.get("reason", "")))
            print("  cleaned turns " + str(start) + "-" + str(start + len(batch) - 1))

        # write cleaned txt
        with open(os.path.join(args.input, name + ".cleaned.txt"), "w", encoding="utf-8") as f:
            for t in turns:
                f.write("[" + t["ts"] + "] " + t["label"] + ": " + t["text"] + "\n\n")
        # write cleaned docx
        doc = Document(); doc.styles['Normal'].font.size = Pt(11)
        h = doc.add_paragraph(); hr = h.add_run(name + " - cleaned")
        hr.bold = True; hr.font.size = Pt(18); hr.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        note = doc.add_paragraph(); nr = note.add_run(
            "AI-cleaned draft. {{word?}} marks an uncertain guess. See "
            + name + ".FLAGS.txt for spots to check against the audio.")
        nr.italic = True; nr.font.size = Pt(9); nr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()
        for t in turns:
            p = doc.add_paragraph()
            lab = p.add_run(t["label"] + ": "); lab.bold = True
            lab.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)
            p.add_run(t["text"])
        doc.save(os.path.join(args.input, name + ".cleaned.docx"))
        # write flags
        with open(os.path.join(args.input, name + ".FLAGS.txt"), "w", encoding="utf-8") as f:
            f.write("Spots to check against the audio for: " + name + "\n\n")
            if not flags:
                f.write("(none flagged - still worth spot-checking key quotes)\n")
            for ts, lab, reason in flags:
                f.write("[" + ts + "] " + lab + " - " + reason + "\n")
        print("  wrote .cleaned.docx / .cleaned.txt / .FLAGS.txt  (" + str(len(flags)) + " flags)")

    print("\nDone.")


if __name__ == "__main__":
    main()
